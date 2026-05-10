"""Generate CLO4 IDS ML assignment report as .docx. Run: python generate_CLO4_report.py"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_title(text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)


def add_center_line(text: str, bold=False, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def fig_placeholder(num: str, caption: str, where: str):
    p = doc.add_paragraph()
    p.add_run(f"[INSERT FIGURE {num} HERE — {where}]").italic = True
    cap = doc.add_paragraph()
    r = cap.add_run(f"Figure {num}: {caption}")
    r.italic = True


# --- Title page ---
add_title("AI-Powered Network Intrusion Detection (Proof of Concept)")
doc.add_paragraph()
add_center_line("Course: Information Security", size=12)
add_center_line("CLO: 4 — Create solutions to real-life scenarios using security-related tools", size=12)
doc.add_paragraph()
add_center_line("Student Name: [YOUR FULL NAME]", bold=True, size=12)
add_center_line("Student ID: [YOUR STUDENT ID]", bold=True, size=12)
doc.add_paragraph()
add_center_line("Submission: Individual Project Deliverable", size=11)
add_center_line("GitHub Repository: [PASTE PUBLIC REPO URL, e.g. https://github.com/you/CLO4-IDS-ML-Solution]", size=10)
doc.add_page_break()

# --- Executive Summary ---
doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    "Organizations cannot rely on firewalls alone; they need intelligent ways to spot malicious "
    "network behaviour in large volumes of flow data. This project delivers a proof-of-concept "
    "machine learning classifier for SecureNet Corp. that labels network flows using the CIC-IDS2017 "
    "family of captures (DDoS, port scan, web attacks, brute-force, infiltration, and related threats). "
    "Raw CSV files were merged (approximately 2.83 million flows, 79 columns including the label), "
    "cleaned, scaled, and split into training and held-out test sets. Two models were trained: a "
    "Random Forest ensemble (scikit-learn) and a deep Convolutional–LSTM model (TensorFlow/Keras). "
    "On the held-out test set, the Random Forest reached about 99.87% overall accuracy, while the "
    "CNN–LSTM reached about 98.69% accuracy. Per-class precision and recall show that most attack "
    "families are detected reliably, but very rare classes (for example some web-injection cases with "
    "almost no test examples) can produce dangerously low recall—meaning some real attacks could be "
    "missed in production. The notebook, metrics, and figures support a management-ready view of "
    "what ML can and cannot promise for intrusion detection."
)

# --- Introduction ---
doc.add_heading("Introduction and Real-World Scenario", level=1)
doc.add_paragraph(
    "SecureNet Corp. operates services that are continuously probed for weaknesses: denial-of-service "
    "floods, port scanning, credential brute-forcing, and web-layer attacks. Traditional rules and "
    "signatures catch known patterns but struggle with volume, polymorphic behaviour, and subtle "
    "blends of benign and attack traffic."
)
doc.add_paragraph(
    "The Chief Information Security Officer (CISO) asked for a feasibility study: can supervised "
    "machine learning augment an existing Network Intrusion Detection System (NIDS) by automatically "
    "classifying flows as benign or as specific attack types? This report documents the end-to-end "
    "workflow—data engineering, model training, evaluation, and security interpretation—mirroring "
    "how a security team would prototype a new analytics tool before any real-time deployment."
)
doc.add_paragraph(
    "Scope note on “real time”: this deliverable implements offline training and batch evaluation on "
    "a labelled dataset. A production “real-time” NIDS would apply the same trained model to each "
    "new flow or window of packets within milliseconds; that deployment step (streaming features, "
    "latency budgets, and integration with SIEM) is proposed under Future Enhancements."
)

# --- Methodology ---
doc.add_heading("Methodology", level=1)

doc.add_heading("Dataset and scenario justification", level=2)
doc.add_paragraph(
    "The Canadian Institute for Cybersecurity Intrusion Detection System 2017 (CIC-IDS2017) dataset "
    "is widely used in research and includes realistic benign background traffic mixed with modern "
    "attacks captured in separate day/time CSV files (e.g. Monday through Friday working hours). "
    "In this project, eight CSV files were combined, including scenarios such as DDoS, PortScan, "
    "Web Attacks, Infiltration, and multiple Denial-of-Service variants, plus Patator-based "
    "brute-force traffic. That diversity matches SecureNet’s concern set: volumetric attacks, "
    "reconnaissance, application-layer abuse, and lateral movement style events."
)
doc.add_paragraph(
    "After merging, the combined table contained about 2,830,743 rows and 79 columns (78 numerical "
    "flow features and one label column). The label distribution is highly imbalanced: benign flows "
    "dominate, while classes such as Infiltration or Heartbleed are extremely rare—an important "
    "reality for both training and evaluation."
)

doc.add_heading("Feature inventory (columns)", level=2)
doc.add_paragraph(
    "The dataset’s columns are flow statistics (e.g. duration, packet counts, byte lengths, IAT "
    "features, TCP flags, window sizes). Capturing the full column list documents traceability for "
    "the CISO and auditors."
)
fig_placeholder(
    "1",
    "Complete list of feature column names (Index object from pandas after stripping whitespace), "
    "showing all flow metrics and the final Label field.",
    "Screenshot from your notebook: output of data.columns (or data.columns.tolist()).",
)

doc.add_heading("Exploratory Data Analysis (EDA)", level=2)
doc.add_paragraph(
    "EDA steps included inspecting the schema, previewing rows with data.head(), and measuring class "
    "frequencies with value_counts() on the Label column. The counts confirm heavy imbalance: benign "
    "traffic forms the majority bucket, while specialised attacks appear in much smaller numbers. "
    "That imbalance explains why overall accuracy can look excellent while individual rare attacks "
    "still score poorly on recall."
)
fig_placeholder(
    "2",
    "Label distribution (value_counts) showing benign vs each attack type.",
    "Screenshot: output of data['Label'].value_counts() from the notebook.",
)

doc.add_heading("Data preprocessing", level=2)
doc.add_paragraph(
    "Preprocessing followed standard practice for numeric flow records:"
)
items = [
    "Infinity cleanup: replace ±inf with NaN, then drop rows containing NaN so models do not see non-finite values.",
    "Target encoding: scikit-learn LabelEncoder mapped each attack name to an integer class index.",
    "Feature scaling: StandardScaler standardized the 78 numeric features to zero mean and unit variance, "
    "so distance- and gradient-based models are not dominated by large-magnitude columns.",
    "Train/test split: 80% train, 20% test with random_state=42 for reproducibility.",
]
for t in items:
    doc.add_paragraph(t, style="List Bullet")

doc.add_paragraph(
    "Note for improvement: in a stricter experimental protocol, the scaler should be fit only on the "
    "training split and then applied to the test split to avoid information leakage; this refinement "
    "should be mentioned in the viva and can be adopted before final submission."
)
doc.add_paragraph(
    "CIC-IDS2017 CSVs used here are predominantly numeric; there were no separate categorical fields "
    "such as “protocol” requiring one-hot encoding beyond the label."
)

doc.add_heading("Machine learning tool design", level=2)
doc.add_paragraph(
    "Two complementary learners were implemented to satisfy the brief (scikit-learn and TensorFlow)."
)

doc.add_heading("Random Forest (scikit-learn)", level=3)
doc.add_paragraph(
    "A RandomForestClassifier with 100 trees was chosen because tree ensembles handle non-linear "
    "relationships among dozens of flow features, are robust to outliers after scaling, provide "
    "feature-importance diagnostics, and often achieve strong accuracy on tabular security data "
    "without extensive hyperparameter tuning. Random Forests are a common baseline in IDS literature "
    "for multiclass flow classification."
)

doc.add_heading("Convolutional 1D + LSTM (TensorFlow/Keras)", level=3)
doc.add_paragraph(
    "A deep model reshaped each sample as a sequence of length 78 with one channel per timestep, then "
    "applied Conv1D layers (local pattern extraction), MaxPooling, LSTM layers (temporal dependency "
    "across the synthetic sequence), Dropout for regularization, and a Dense softmax output for "
    "multiclass prediction. This satisfies the TensorFlow requirement and explores whether sequence "
    "structure helps compared to the tree ensemble."
)

# --- Results ---
doc.add_heading("Results and Security Analysis", level=1)
doc.add_paragraph(
    "Both models were evaluated on the same held-out test partition (565,576 samples). Overall "
    "accuracy is high for both approaches, but security analysts must prioritise per-class recall "
    "for attacks—missing an attack (false negative) is often worse than raising an extra alert "
    "(false positive), depending on SOC policy."
)

doc.add_heading("Overall accuracy", level=2)
doc.add_paragraph(
    "Random Forest test accuracy: approximately 0.9987 (99.87%). "
    "CNN–LSTM test accuracy: approximately 0.9869 (98.69%). "
    "The ensemble slightly outperformed the deep model on this split, which is plausible given the "
    "tabular nature of the features and class imbalance."
)

doc.add_heading("Precision, recall, and confusion matrix (Random Forest)", level=2)
doc.add_paragraph(
    "The sklearn classification_report on the test set provides precision, recall, and F1 per class. "
    "Illustrative findings from the saved notebook run:"
)
analysis_bullets = [
    "Benign traffic (majority class): precision and recall both about 1.00 on hundreds of thousands of test samples.",
    "Bot traffic: precision ≈ 0.86, recall ≈ 0.75 — about one quarter of bot flows could be missed, "
    "which in operations might mean some command-and-control or botnet behaviour slips through unless "
    "complemented with other controls.",
    "Several high-volume attack families (e.g. DDoS, DoS variants, PortScan) show near-perfect precision and recall.",
    "Rare web-attack and injection-related classes: some show recall near 0.36 or even 0.00 when only "
    "one or a handful of test samples exist — statistically unstable and operationally dangerous if "
    "interpreted as “solved” based on accuracy alone.",
]
for t in analysis_bullets:
    doc.add_paragraph(t, style="List Bullet")

fig_placeholder(
    "3",
    "Random Forest confusion matrix heatmap (test set).",
    "Screenshot: sns.heatmap of confusion_matrix(y_test, y_pred) with title “Random Forest Confusion Matrix”.",
)

fig_placeholder(
    "4",
    "Random Forest: full classification_report output (precision, recall, F1 per class).",
    "Screenshot: notebook cell printing classification_report(y_test, y_pred).",
)

doc.add_heading("Deep learning training curve", level=2)
fig_placeholder(
    "5",
    "Training vs validation accuracy across epochs for the CNN–LSTM model.",
    "Screenshot: plot of history.history['accuracy'] and val_accuracy.",
)

doc.add_heading("Model comparison", level=2)
doc.add_paragraph(
    "A small summary table (as in the notebook) contrasts Random Forest vs CNN–LSTM on test accuracy. "
    "For management: report both overall metrics and the weakest per-class recalls to avoid over-stating "
    "readiness for deployment."
)
fig_placeholder(
    "6",
    "Side-by-side model comparison (pandas DataFrame: Model vs Accuracy).",
    "Screenshot: the comparison DataFrame from the notebook.",
)

doc.add_heading("Security interpretation (critical analysis)", level=2)
doc.add_paragraph(
    "High overall accuracy is misleading when the dataset is dominated by benign flows. A model could "
    "score 99%+ while still missing entire attack families that represent critical risk. Low recall "
    "on Bot or specific Web Attack labels means many malicious sessions would be classified as benign—"
    "the digital equivalent of silent failure for an IDS. Low precision on a class would mean more "
    "false alarms for that label, increasing analyst fatigue."
)
doc.add_paragraph(
    "Mitigations aligned with security operations include: stratified sampling or class weighting, "
    "cost-sensitive learning, collecting more real examples of rare attacks, anomaly detection for "
    "unknown classes, and always pairing ML scores with rule-based and threat-intelligence context."
)

# --- Conclusion ---
doc.add_heading("Conclusion and Future Enhancements", level=1)
doc.add_paragraph(
    "This project demonstrates a complete ML lifecycle for multiclass network intrusion detection using "
    "CIC-IDS2017-style data: ingestion, cleaning, encoding, scaling, training with both Random Forest "
    "and a TensorFlow CNN–LSTM architecture, and evaluation with accuracy, per-class precision/recall, "
    "and confusion matrices. The results support the CISO’s question—ML can strongly augment a NIDS—"
    "but they also show that rare attacks and class imbalance demand careful metrics and operational "
    "design, not headline accuracy alone."
)
doc.add_paragraph("Future enhancements:")
for t in [
    "Fit preprocessing only on training data; validate with cross-validation and stratified splits.",
    "Address imbalance: class weights, resampling, or focal loss in the neural model.",
    "Binary “attack vs benign” model alongside multiclass for triage.",
    "Feature importance (Random Forest) and SHAP-style explanations for incident reviewers.",
    "Real-time path: stream features from Zeek/Suricata, enforce latency SLOs, and log model version.",
    "Continuous retraining and adversarial robustness testing.",
]:
    doc.add_paragraph(t, style="List Bullet")

# --- References ---
doc.add_heading("References", level=1)
refs = [
    "Sharafaldin, I., Habibi Lashkari, A., & Ghorbani, A. A. (2018). Toward generating a new intrusion detection dataset and intrusion traffic characterization. ICISSp.",
    "Scikit-learn developers. Machine learning in Python (documentation for RandomForestClassifier, preprocessing, metrics). https://scikit-learn.org/stable/",
    "TensorFlow team. Keras API documentation (Conv1D, LSTM, Sequential). https://www.tensorflow.org/",
    "Canadian Institute for Cybersecurity. CIC-IDS2017 dataset description and download pages (use the citation/version you actually retrieved).",
]
for r in refs:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("GitHub profile link", level=1)
doc.add_paragraph("[Paste your GitHub profile URL here, e.g. https://github.com/yourusername]")

out = r"c:\Users\maazs\OneDrive\Desktop\Assignment\CLO4_IDS_ML_Report.docx"
doc.save(out)
print("Saved:", out)
